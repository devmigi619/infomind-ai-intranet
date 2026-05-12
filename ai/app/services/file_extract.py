"""파일 확장자별 텍스트 추출.

지원: pdf, docx, xlsx, txt, md, csv. 그 외는 빈 문자열 반환.
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """확장자를 감지해 적절한 추출기로 텍스트를 뽑는다.

    추출 실패 시 빈 문자열 반환 (호출자가 skip할 수 있도록).
    """
    ext = _get_ext(filename)

    try:
        if ext == "pdf":
            return _extract_pdf(file_bytes)
        if ext == "docx":
            return _extract_docx(file_bytes)
        if ext == "xlsx":
            return _extract_xlsx(file_bytes)
        if ext in ("txt", "md", "csv"):
            return _decode_text(file_bytes)
    except Exception as e:
        logger.warning("파일 텍스트 추출 실패 (filename=%s, ext=%s): %s", filename, ext, e)
        return ""

    logger.info("지원하지 않는 확장자: %s (filename=%s)", ext, filename)
    return ""


def _get_ext(filename: str) -> Optional[str]:
    if not filename or "." not in filename:
        return None
    return filename.rsplit(".", 1)[1].lower()


def _extract_pdf(data: bytes) -> str:
    import fitz  # pymupdf

    text_parts: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # 표 셀도 포함
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"[{sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _decode_text(data: bytes) -> str:
    # utf-8 우선, 실패 시 cp949(국내 텍스트 자주 발생) 시도
    for enc in ("utf-8", "cp949", "latin-1"):
        try:
            return data.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return ""
